#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Word文档生成器：负责将题目内容转换为Word文档
"""

import os
import logging
import requests
import tempfile
from io import BytesIO
from typing import Dict, Any, List, Union
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class WordGenerator:
    """Word文档生成器类，负责创建和格式化Word文档"""
    
    def __init__(self, template_path=None):
        """
        初始化Word文档生成器
        
        Args:
            template_path: Word模板文件路径，如果为None则创建新文档
        """
        self.template_path = template_path
        logger.info("初始化Word文档生成器")
    
    def create_document(self, content: List[Dict], output_path: str) -> str:
        """
        创建Word文档
        
        Args:
            content: 文档内容，包含题目信息的列表
            output_path: 输出文件路径
            
        Returns:
            生成的文档路径
        """
        logger.info(f"开始创建Word文档，输出路径: {output_path}")
        
        # 创建文档对象
        if self.template_path and os.path.exists(self.template_path):
            doc = Document(self.template_path)
            logger.info(f"使用模板创建文档: {self.template_path}")
        else:
            doc = Document()
            logger.info("创建新文档")
        
        # 设置文档样式
        self._setup_document_styles(doc)
        
        # 添加文档标题
        self._add_title(doc, "小学二年级仿写练习题与看图写话练习题")
        
        # 收集所有参考答案，按大题分组
        answers = []
        
        # 处理每种题型
        for item in content:
            if "maxTitle" in item and "questions" in item:
                # 添加大题标题
                self._add_section_title(doc, item["maxTitle"])
                
                # 添加题目要求
                if "require" in item:
                    self._add_requirement(doc, item["require"])
                
                # 收集当前大题的答案
                current_answers = {"title": item["maxTitle"], "answers": []}
                
                # 处理题目
                if item["maxTitle"] == "句子仿写练习题":
                    self._add_imitation_questions(doc, item["questions"], current_answers)
                elif item["maxTitle"] == "看图写话练习题":
                    self._add_picture_questions(doc, item["questions"], current_answers)
                
                # 添加到总答案列表
                if current_answers["answers"]:
                    answers.append(current_answers)
        
        # 添加分页符
        doc.add_page_break()
        
        # 添加答案部分
        self._add_answers_section(doc, answers)
        
        # 保存文档
        doc.save(output_path)
        
        logger.info(f"文档创建完成: {output_path}")
        return output_path
    
    def _setup_document_styles(self, doc):
        """设置文档样式"""
        # 设置中文字体
        style_names = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']
        for style_name in style_names:
            style = doc.styles[style_name]
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            # 设置中文字体
            font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 创建自定义样式
        if 'Title' not in doc.styles:
            title_style = doc.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(18)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            # 设置中文字体
            title_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        if 'Section' not in doc.styles:
            section_style = doc.styles.add_style('Section', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Times New Roman'
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(8)
            # 设置中文字体
            section_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        if 'Requirement' not in doc.styles:
            req_style = doc.styles.add_style('Requirement', WD_STYLE_TYPE.PARAGRAPH)
            req_style.font.name = 'Times New Roman'
            req_style.font.size = Pt(14)
            req_style.font.italic = True
            req_style.paragraph_format.space_after = Pt(8)
            # 设置中文字体
            req_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        if 'Question' not in doc.styles:
            q_style = doc.styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
            q_style.font.name = 'Times New Roman'
            q_style.font.size = Pt(12)
            q_style.paragraph_format.space_after = Pt(6)
            q_style.paragraph_format.first_line_indent = Inches(0.25)
            # 设置中文字体
            q_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        if 'Answer' not in doc.styles:
            a_style = doc.styles.add_style('Answer', WD_STYLE_TYPE.PARAGRAPH)
            a_style.font.name = 'Times New Roman'
            a_style.font.size = Pt(12)
            a_style.font.italic = True
            a_style.paragraph_format.space_after = Pt(6)
            a_style.paragraph_format.left_indent = Inches(0.5)
            # 设置中文字体
            a_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _add_title(self, doc, title):
        """添加文档标题"""
        paragraph = doc.add_paragraph(title, style='Title')
        doc.add_paragraph()  # 添加空行
    
    def _add_section_title(self, doc, title):
        """添加大题标题"""
        paragraph = doc.add_paragraph(title, style='Section')
    
    def _add_requirement(self, doc, requirement):
        """添加题目要求"""
        paragraph = doc.add_paragraph(requirement, style='Requirement')
    
    def _download_image(self, url):
        """
        从URL下载图片
        
        Args:
            url: 图片URL
            
        Returns:
            图片的二进制数据
        """
        try:
            response = requests.get(url, timeout=30)
            if response.status_code == 200:
                return BytesIO(response.content)
            else:
                logger.error(f"下载图片失败，状态码: {response.status_code}")
                return None
        except Exception as e:
            logger.error(f"下载图片时发生错误: {str(e)}")
            return None
    
    def _add_imitation_questions(self, doc, questions, answers_collector):
        """添加句子仿写题"""
        for question in questions:
            # 添加题号和问题
            q_text = f"{question['number']}. {question['question']}"
            paragraph = doc.add_paragraph(q_text, style='Question')
            
            # 添加仿写空格
            if "Imitate_writing" in question and isinstance(question["Imitate_writing"], list):
                imitate_text = ""
                for i, word in enumerate(question["Imitate_writing"]):
                    if i > 0:
                        # 添加下划线空格
                        imitate_text += "_____________"
                    imitate_text += word
                
                # 最后一个词后也添加下划线
                imitate_text += "_____________"
                
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(imitate_text)
            
            # 收集参考答案
            if "reference_answer" in question:
                answers_collector["answers"].append({
                    "number": question["number"],
                    "question": question["question"],
                    "answer": question["reference_answer"]
                })
            
            # 添加空行
            doc.add_paragraph()
    
    def _add_picture_questions(self, doc, questions, answers_collector):
        """添加看图写话题"""
        for question in questions:
            # 添加题号和问题
            q_text = f"{question['number']}. {question['question']}"
            paragraph = doc.add_paragraph(q_text, style='Question')
            
            # 添加图片 - 优先使用本地图片路径，如果没有则尝试从URL下载
            image_added = False
            
            # 如果有本地图片路径，直接使用本地图片
            if "local_image_path" in question and question["local_image_path"]:
                try:
                    local_path = question["local_image_path"]
                    if os.path.exists(local_path):
                        # 添加图片到文档
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(local_path, width=Inches(5))
                        logger.info(f"成功从本地添加图片到题目 {question['number']}")
                        image_added = True
                    else:
                        logger.warning(f"本地图片文件不存在: {local_path}")
                except Exception as e:
                    logger.error(f"添加本地图片时发生错误: {str(e)}")
            
            # 如果没有成功添加本地图片，尝试从URL下载
            if not image_added and "image_url" in question and question["image_url"]:
                try:
                    # 下载图片
                    image_data = self._download_image(question["image_url"])
                    if image_data:
                        # 添加图片到文档
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(image_data, width=Inches(5))
                        logger.info(f"成功从URL添加图片到题目 {question['number']}")
                        image_added = True
                    else:
                        logger.warning(f"无法下载图片，题目 {question['number']}")
                except Exception as e:
                    logger.error(f"添加URL图片时发生错误: {str(e)}")
            
            # 不再在图片下方显示图片描述，而是将其添加到参考答案中
            
            # 添加写作空间
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run("写话：").bold = True
            
            # 添加多行空白供学生写作
            for i in range(8):
                doc.add_paragraph("_" * 60, style='Normal')
                # 在循环最后一次加上句号
                if i == 7:
                    doc.add_paragraph("。", style='Normal')
            
            # 收集参考答案和图片描述
            if "reference_answer" in question:
                answer_data = {
                    "number": question["number"],
                    "question": question["question"],
                    "answer": question["reference_answer"]
                }
                
                # 如果有图片描述，添加到参考答案中
                if "prompt" in question:
                    answer_data["prompt"] = question["prompt"]
                
                answers_collector["answers"].append(answer_data)
            
            # 添加空行
            doc.add_paragraph()
    
    def _add_answers_section(self, doc, answers):
        """添加答案部分"""
        # 添加答案标题
        self._add_title(doc, "参考答案")
        
        # 处理每种题型的答案
        for answer_group in answers:
            # 添加大题标题
            self._add_section_title(doc, f"{answer_group['title']}参考答案")
            
            # 添加每道题的答案
            for answer_item in answer_group["answers"]:
                # 添加题号和问题
                q_text = f"{answer_item['number']}. {answer_item['question']}"
                paragraph = doc.add_paragraph(q_text, style='Question')
                
                # 如果是看图写话题目且有图片描述，先添加图片描述
                if answer_group['title'] == "看图写话练习题" and "prompt" in answer_item:
                    p = doc.add_paragraph(style='Answer')
                    run = p.add_run(f"图片描述：{answer_item['prompt']}")
                    run.italic = True
                
                # 添加参考答案
                p = doc.add_paragraph(style='Answer')
                run = p.add_run(f"参考答案：{answer_item['answer']}")
                run.italic = True
            
            # 在不同大题之间添加空行
            doc.add_paragraph()


def main():
    """测试函数"""
    # 简单的测试数据
    test_content = [
            {
                "maxTitle": "句子仿写练习题",
                "require": "请根据例句进行仿写。",
                "questions": [
                    {
                        "number": 1,
                        "question": "例句：小鸟在树上唱歌。",
                        "reference_answer": "小鸭在河里游泳。",
                        "Imitate_writing": [
                            "小鸭在",
                            "",
                            "里",
                            ""
                        ]
                    },
                    {
                        "number": 2,
                        "question": "例句：小狗在院子里玩耍。",
                        "reference_answer": "小猫在窗台上晒太阳。",
                        "Imitate_writing": [
                            "小猫在",
                            "",
                            "上",
                            ""
                        ]
                    }
                ]
            },
            {
                "maxTitle": "看图写话练习题",
                "require": "请根据图片描述进行写话。",
                "questions": [
                    {
                        "number": 1,
                        "question": "请描述图片中的场景。",
                        "prompt": "图片中是一个阳光明媚的公园，草地上有许多小朋友在玩耍。一个小男孩正在放风筝，风筝在天空中高高飘扬。一个小女孩在草地上追逐着一只小狗，小狗的尾巴摇得很快。远处有一棵大树，树下有一位老爷爷正在看报纸。天空中飘着几朵白云，整个场景充满了欢乐和宁静。",
                        "reference_answer": "图片中是一个阳光明媚的公园，草地上有许多小朋友在玩耍。一个小男孩正在放风筝，风筝在天空中高高飘扬。一个小女孩在草地上追逐着一只小狗，小狗的尾巴摇得很快。远处有一棵大树，树下有一位老爷爷正在看报纸。天空中飘着几朵白云，整个场景充满了欢乐和宁静。"
                    }
                ]
            }
        ]
    
    # 创建Word文档
    generator = WordGenerator()
    output_path = "test_exercises.docx"
    generator.create_document(test_content, output_path)
    print(f"测试文档已创建: {output_path}")


if __name__ == "__main__":
    main() 